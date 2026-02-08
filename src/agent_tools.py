from langchain_core.tools import Tool, tool
from langchain_community.tools import DuckDuckGoSearchRun
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

# --- TOOL 1: THE SMART WORD PROCESSOR ---
@tool
def create_word_document(content: str):
    """
    Useful for creating Word documents. 
    Input should be the raw text content.
    The tool AUTOMATICALLY forces the current date and handles formatting.
    """
    try:
        doc = Document()
        
        # 1. Global Font Settings
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.space_after = Pt(0) 

        # 2. DATE FORCE-FIX (The "System Override")
        # Get the real current date
        real_today_str = datetime.now().strftime("%B %d, %Y")
        
        # Split lines to process them
        raw_lines = content.split('\n')
        clean_lines = []
        
        # We filter the agent's output. 
        # If the agent tried to write "Date: Feb 20...", we DELETE it.
        for line in raw_lines:
            # Check if line starts with "Date:" or matches a date-like pattern in the header
            if line.strip().lower().startswith("date:"):
                continue # Skip this line, it's likely wrong
            clean_lines.append(line)
            
        # NOW, we insert the CORRECT date at the very top
        clean_lines.insert(0, f"Date: {real_today_str}")
        
        # 3. DOCUMENT GENERATION LOOP
        header_paragraph = None
        
        for line in clean_lines:
            clean_line = line.strip()
            lower_line = clean_line.lower()

            if not clean_line:
                continue

            # Remove artifacts
            if "memorandum" in lower_line or "dress code:" == lower_line:
                continue

            # --- HEADER LOGIC (Date, To, From, Re) ---
            if lower_line.startswith(("to:", "from:", "date:", "re:")) or lower_line.startswith(("to.", "from.")):
                
                if header_paragraph is None:
                    header_paragraph = doc.add_paragraph()
                    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    header_paragraph.add_run().add_break()
                
                # Handling "Key: Value" formatting
                separator = ":" if ":" in clean_line else "."
                if separator in clean_line:
                    parts = clean_line.split(separator, 1)
                    if len(parts) > 1:
                        run_key = header_paragraph.add_run(parts[0] + ":")
                        run_key.bold = True
                        run_val = header_paragraph.add_run(parts[1])
                    else:
                        header_paragraph.add_run(clean_line)
                else:
                    run = header_paragraph.add_run(clean_line)
                    if lower_line.startswith("date:"):
                         pass 
                
                continue

            header_paragraph = None 

            # --- SUBJECT LOGIC ---
            if lower_line.startswith("subject:"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                
                run = p.add_run(clean_line)
                run.bold = True
                continue

            # --- SIGNATURE LOGIC ---
            is_signature = False
            if lower_line.startswith(("sincerely", "regards", "best,")):
                is_signature = True
            elif "mark" in lower_line and "ai department" in lower_line:
                is_signature = True
            
            p = doc.add_paragraph()
            if is_signature:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(12)

            # Bold markdown support
            parts = clean_line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1: run.bold = True

        # Save
        current_dir = os.path.dirname(os.path.abspath(__file__))
        save_dir = os.path.join(os.path.dirname(current_dir), "data")
        
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)
            
        filename = os.path.join(save_dir, "agent_output_final.docx")
        doc.save(filename)
        return f"Document created successfully: {filename}"

    except Exception as e:
        return f"Error creating document: {str(e)}"

# --- TOOL 2: WEB SEARCH ---
def get_search_tool():
    return Tool(
        name="web_search",
        func=DuckDuckGoSearchRun().run,
        description="Useful for searching the internet."
    )

# --- TOOL 3: DATE GETTER ---
@tool
def get_todays_date(input_str: str = ""):
    """
    Returns the actual current date from the computer's internal clock.
    """
    return datetime.now().strftime("%B %d, %Y")